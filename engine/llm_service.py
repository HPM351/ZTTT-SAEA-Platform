# engine/llm_service.py
import os
import json
import io
import httpx
from fastapi import APIRouter, Depends, File, Form, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from engine.api_chat_history import get_current_user

llm_router = APIRouter(prefix="/api/llm", tags=["LLM Assistant"])

# pypdf — PyPDF2 的继任者，API 兼容
from pypdf import PdfReader
from pypdf.errors import PdfReadWarning
import warnings

# 屏蔽不规范 PDF 的冗余警告，避免中断解析流程
warnings.filterwarnings("ignore", category=PdfReadWarning)

try:
    import docx
except ImportError:
    docx = None

# olefile — 解析旧版 Word .doc 二进制格式 (OLE2 容器)
try:
    import olefile
except ImportError:
    olefile = None

# 限制常量
MAX_FILE_SIZE = 20 * 1024 * 1024   # 20 MB
MAX_PDF_PAGES = 15
MAX_EXTRACTED_CHARS = 50000  # 所有附件累计提取文本上限

# 已初始化 system prompt 的 session 集合（重启后重置，但网关 session transcript 保留了 soul）
_init_sessions = set()


class VisionChatRequest(BaseModel):
    message: str
    history: list = []
    images_base64_list: list[str] = []


# ==========================================
# 1. 文献助手专属路由 (纯文本) — multipart/form-data
# ==========================================
@llm_router.post("/chat/text")
async def chat_with_text_model(
    message: str = Form(...),
    session_id: str = Form("default_session"),
    files: list[UploadFile] = File(default=[]),
    current_user = Depends(get_current_user),
):
    print(f"🚀 收到来自研究员 [{current_user.username}] 的文献对话请求 (Session: {session_id})")

    OPENCLAW_URL = "http://127.0.0.1:18789/v1/chat/completions"
    OPENCLAW_TOKEN = "123456"
    target_model = "openclaw:main"

    # ----------------------------------------------------------------
    # 2. 加载知识库 Prompt（纯静态 soul）
    # ----------------------------------------------------------------
    base_soul = "你是一个专业的高功率微波与计算机科学文献助手。请基于专业知识准确回答。"
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    prompt_path = os.path.join(base_dir, "prompts", "saea_main.md")
    if os.path.exists(prompt_path):
        with open(prompt_path, "r", encoding="utf-8") as f:
            base_soul = f.read()

    # ----------------------------------------------------------------
    # 3. 逐文件解析
    # ----------------------------------------------------------------
    print(f"📎 附件数量: {len(files)}")
    for f_tmp in files:
        print(f"  - filename='{f_tmp.filename}', size={f_tmp.size}")
    extracted_text = ""
    for file in files:
        if len(extracted_text) >= MAX_EXTRACTED_CHARS:
            print(f"⏭️ 累计提取文本已达 {MAX_EXTRACTED_CHARS} 字上限，跳过剩余文件")
            break
        file_name = file.filename or "未命名文件"
        try:
            content_bytes = await file.read()
        except Exception as e:
            print(f"读取 {file_name} 失败: {e}")
            continue

        if len(content_bytes) > MAX_FILE_SIZE:
            extracted_text += (
                f"\n\n--- 附件 [{file_name}] --- (文件超过 {MAX_FILE_SIZE // 1024 // 1024}MB 限制，已跳过)\n"
            )
            continue

        lower_name = file_name.lower()

        if lower_name.endswith(".txt") or lower_name.endswith(".md"):
            extracted_text += f"\n\n--- 附件 [{file_name}] 内容开始 ---\n"
            extracted_text += content_bytes.decode("utf-8", errors="ignore")
            extracted_text += f"\n--- 附件 [{file_name}] 内容结束 ---\n"

        elif lower_name.endswith(".pdf"):
            extracted_text += f"\n\n--- 附件 [{file_name}] 内容开始 ---\n"
            try:
                pdf_reader = PdfReader(io.BytesIO(content_bytes))
                for i, page in enumerate(pdf_reader.pages):
                    if i >= MAX_PDF_PAGES:
                        extracted_text += f"[截断：文件过长，仅提取前 {MAX_PDF_PAGES} 页]\n"
                        break
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
            except Exception as e:
                extracted_text += f"[解析失败: {e}]\n"
            extracted_text += f"\n--- 附件 [{file_name}] 内容结束 ---\n"

        elif lower_name.endswith(".docx"):
            extracted_text += f"\n\n--- 附件 [{file_name}] 内容开始 ---\n"
            if docx:
                try:
                    doc = docx.Document(io.BytesIO(content_bytes))
                    for para in doc.paragraphs:
                        if para.text.strip():
                            extracted_text += para.text + "\n"
                    for table in doc.tables:
                        row_data = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                        extracted_text += "\n".join(row_data) + "\n"
                except Exception as e:
                    extracted_text += f"[解析失败: {e}]\n"
            else:
                extracted_text += "[提示：后端未安装 python-docx 库]\n"
            extracted_text += f"\n--- 附件 [{file_name}] 内容结束 ---\n"

        elif lower_name.endswith(".doc") and not lower_name.endswith(".docx"):
            extracted_text += f"\n\n--- 附件 [{file_name}] 内容开始 ---\n"
            if olefile:
                try:
                    ole = olefile.OleFileIO(io.BytesIO(content_bytes))
                    if ole.exists("WordDocument"):
                        data = ole.openstream("WordDocument").read()
                        chars = []
                        i = 0
                        while i < len(data) - 1:
                            if data[i+1] == 0x00 and 0x20 <= data[i] <= 0x7E:
                                chars.append(chr(data[i]))
                            i += 2
                        raw_text = "".join(chars)
                        if raw_text.strip():
                            extracted_text += raw_text + "\n"
                        else:
                            extracted_text += "[解析完成，但未提取到可读文本（可能是加密文档）]\n"
                    else:
                        extracted_text += "[无法识别为有效的 Word .doc 格式]\n"
                    ole.close()
                except Exception as e:
                    extracted_text += f"[解析失败: {e}]\n"
            else:
                extracted_text += "[提示：后端未安装 olefile 库]\n"
            extracted_text += f"\n--- 附件 [{file_name}] 内容结束 ---\n"

        else:
            extracted_text += f"\n\n--- 附件 [{file_name}] --- (暂不支持此格式)\n"

    # ----------------------------------------------------------------
    # 4. 组装最终用户消息
    # ----------------------------------------------------------------
    final_user_message = message
    if extracted_text:
        final_user_message += f"\n\n以下是用户提供的参考资料，请结合上述问题和资料进行回答：{extracted_text}"
        print(f"📄 提取文本长度: {len(extracted_text)} 字符")

    # ----------------------------------------------------------------
    # 5. 组装消息——走网关 SESSION 管理上下文
    #    第一轮带 system prompt 初始化 session，
    #    后续轮次只发当前用户消息，网关 session 自动继承历史。
    #    无需后端读 DB、无需截断，DeepSeek 原生 1M 上下文全用。
    # ----------------------------------------------------------------
    messages = []
    if session_id not in _init_sessions:
        # 新 session：带 system prompt 初始化
        messages.append({"role": "system", "content": base_soul})
        _init_sessions.add(session_id)
        print(f"🆕 新 session [{session_id}]，初始化 system prompt")
    # 始终只发当前用户消息
    messages.append({"role": "user", "content": final_user_message})

    print(f"📤 发送消息数: {len(messages)}，长度: {len(final_user_message)} 字符")

    payload = {
        "model": target_model,
        "messages": messages,
        "temperature": 0.7,
        "stream": True,
        "user": session_id,
    }

    headers = {
        "Authorization": f"Bearer {OPENCLAW_TOKEN}",
        "Content-Type": "application/json",
    }

    async def stream_generator():
        try:
            async with httpx.AsyncClient() as client:
                async with client.stream(
                    "POST", OPENCLAW_URL, json=payload, headers=headers, timeout=180.0
                ) as response:
                    if response.status_code != 200:
                        error_detail = await response.aread()
                        yield f"data: {{\"error\": \"API响应错误: {error_detail.decode()}\"}}\n\n".encode("utf-8")
                        return
                    async for chunk in response.aiter_bytes():
                        yield chunk
        except Exception as e:
            yield f"data: {{\"error\": \"后端文本路由异常: {str(e)}\"}}\n\n".encode("utf-8")

    return StreamingResponse(stream_generator(), media_type="text/event-stream")


# ==========================================
# 2. 视觉核心专属路由 (多模态)
# ==========================================
@llm_router.post("/chat/vision")
async def chat_with_vision_model(req: VisionChatRequest):
    SILICONFLOW_URL = "https://api.siliconflow.cn/v1/chat/completions"
    SILICONFLOW_KEY = "sk-ypnnksenlishqzdhkrcoiqrgmrjsjqibkawsbpxcnkywxbhp"
    VISION_MODEL = "Qwen/Qwen3.5-397B-A17B"

    saea_soul = "你是一个 SAEA 平台视觉助手，专门用于分析图表和物理曲线。"
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    prompt_path = os.path.join(base_dir, "prompts", "saea_vision.md")
    if os.path.exists(prompt_path):
        with open(prompt_path, "r", encoding="utf-8") as f:
            saea_soul = f.read()

    user_content = []
    for img_b64 in req.images_base64_list:
        user_content.append({"type": "image_url", "image_url": {"url": img_b64}})
    user_content.append({"type": "text", "text": req.message or "请分析提供的图像。"})

    messages = [{"role": "system", "content": saea_soul}] + req.history + [{"role": "user", "content": user_content}]

    payload = {
        "model": VISION_MODEL,
        "messages": messages,
        "temperature": 0.7,
        "stream": True
    }

    headers = {
        "Authorization": f"Bearer {SILICONFLOW_KEY}",
        "Content-Type": "application/json"
    }

    async def stream_generator():
        try:
            async with httpx.AsyncClient() as client:
                async with client.stream("POST", SILICONFLOW_URL, json=payload, headers=headers, timeout=180.0) as response:
                    if response.status_code != 200:
                        error_detail = await response.aread()
                        yield f"data: {{\"error\": \"API响应错误: {error_detail.decode()}\"}}\n\n".encode("utf-8")
                        return
                    async for chunk in response.aiter_bytes():
                        yield chunk
        except Exception as e:
            yield f"data: {{\"error\": \"后端视觉路由异常: {str(e)}\"}}\n\n".encode("utf-8")

    return StreamingResponse(stream_generator(), media_type="text/event-stream")
